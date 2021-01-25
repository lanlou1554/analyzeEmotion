# coding = Utf-8
import sys
from bs4 import BeautifulSoup
import re
import urllib.request, urllib.error
import xlwt
import sqlite3
from selenium import webdriver
import time
from docx import Document
from docx.oxml.ns import qn
import math


def main(stage1,stage2,stage3,stage4):
    saveOneStage(stage2,2)
    saveOneStage(stage3, 3)
    saveOneStage(stage4, 4)

def saveOneStage(stageInfo,stageNum):
    if stageNum == 2:
        start = 13
    else:
        start = 0
    for i in range(start,len(stageInfo)):
        saveOneUrl(stageInfo[i],i,stageNum)

def saveOneUrl(url,i,stageNum):
    questionInfo, questionAnswers,flag = getData(url)
    if flag:
        savepath ="D:\study\zhuhu\data\stage" + str(stageNum) + "\问题" + str(i) + ".docx"
        saveData(questionInfo, questionAnswers, savepath, url)
    else:
        print("阶段"+str(stageNum)+" 问题"+str(i)+"错误！链接是："+str(url))

sleeptime = 3
minLikeTime = 30
percent = 0.2
findHead = re.compile(r'<h1 class="QuestionHeader-title">(.*?)</h1>',re.S)
findDescription = re.compile(r'<span class="RichText ztext" itemprop="text"><p>(.*?)</p></span></div>',re.S)
findQuestionAnswers = re.compile(r'<span>(.*?) 个回答</span></h4>')
findGoodQuestion = re.compile(r'fill-rule="evenodd"></path></svg></span>好问题 (.*?)</button>')
findQuestionComment = re.compile(r'fill-rule="evenodd"></path></svg></span>(.*?) 条评论</button></div>')
findAnswerContent = re.compile(r'<span class="RichText ztext CopyrightRichText-richText" itemprop="text">(.*?)</span>',re.S)
findLikeNum = re.compile(r'fill-rule="evenodd"></path></svg></span>赞同 (.*?)</button>')
findCommentNum = re.compile(r'fill-rule="evenodd"></path></svg></span>(.*?)</button>')

def convertNumStrToInt(s):
    s =s.replace("\n","").replace(" ","")
    s = s.replace("<!---->","")
    if "万" in s:
        s = int(float(s[0:len(s) - 1]) * 10000)
    return int(s)

def getData(url):
    html,flag = askURL(url)

    soup = BeautifulSoup(html,"html.parser")
    questionInfo = []
    questionAnswers = []
    if not flag:
        return questionInfo,questionAnswers,flag
    headItem = str(soup.find_all('div',class_="QuestionHeader")[0])
    head = re.findall(findHead,headItem)[0]
    try:
        description = re.findall(findDescription,headItem)[0]
    except Exception as e:
        description = ""
    try:
        infoGoodQuestionNum = str(soup.find_all('div',class_="GoodQuestionAction")[0])
    except Exception as e:
        infoGoodQuestionNum = ""
    try:
        infoQuestionCommentNum = str(soup.find_all('div', class_="QuestionHeader-Comment")[0])
    except Exception as e:
        infoQuestionCommentNum = ""
    try:
        infoQuestionAnswerNum = str(soup.find_all('h4', class_="List-headerText")[0])
    except Exception as e:
        infoQuestionAnswerNum = ""
    try:
        goodQuestionNum = convertNumStrToInt(re.findall(findGoodQuestion, infoGoodQuestionNum)[0])
    except Exception as e:
        goodQuestionNum = ""
    try:
        questionCommentNum = convertNumStrToInt(re.findall(findQuestionComment, infoQuestionCommentNum)[0])
    except Exception as e:
        questionCommentNum = ""
    try:
        questionAnswerNum = convertNumStrToInt(re.findall(findQuestionAnswers, infoQuestionAnswerNum)[0])
    except Exception as e:
        questionAnswerNum = ""

    questionInfo.append(head)
    questionInfo.append(description)
    questionInfo.append(goodQuestionNum)
    questionInfo.append(questionCommentNum)
    questionInfo.append(questionAnswerNum)
    flag = True

    try:
        answersItem = soup.find_all('div',class_="List-item")
        for item in answersItem:
            item = str(item)
            answer = []
            delete1 = re.compile(r'(<figure.*?</figure>)',re.S)
            content = re.findall(findAnswerContent,item)[0]
            content = re.sub(delete1,"",content)
            try:
                likeNum = convertNumStrToInt(re.findall(findLikeNum,item)[0])
            except Exception as e:
                likeNum = 0
            temp = re.findall(findCommentNum,item)
            for i in range(len(temp)):
                try:
                    commentNum = convertNumStrToInt(re.findall(findCommentNum, item)[i].replace(" 条评论", ""))
                except Exception as e:
                    continue
            answer.append(content)
            answer.append(likeNum)
            answer.append(commentNum)
            questionAnswers.append(answer)
    except Exception as e:
        flag = False

    return questionInfo,questionAnswers,flag


# 得到制定一个URL的网页内容
def askURL(url):
    flag = True
    pageSourceThree = ""
    try:
        path = "C:/Program Files (x86)/Google/Chrome/Application/chromedriver.exe"
        # option = webdriver.ChromeOptions()
        # option.add_experimental_option('excludeSwitches', ['enable-automation'])
        # driver = webdriver.Chrome(options=option)
        # The homepage link for one upper whose uid is 10330740
        driver = webdriver.Chrome(executable_path=path)
        driver.get(url)
        # Load the url
        try:
            time.sleep(sleeptime)
            driver.find_element_by_class_name("Button.Modal-closeButton.Button--plain").click()
            time.sleep(sleeptime)
            driver.find_element_by_class_name("Button.QuestionRichText-more.Button--plain").click()
            time.sleep(sleeptime)
        except Exception as e:
            i=1

        flag = False
        oneFlag = True
        while True:
            driver.execute_script('window.scrollBy(0,document.body.scrollHeight)')
            if oneFlag:
                driver.execute_script('window.scrollBy(document.body.scrollHeight,document.body.scrollHeight-2)')
                oneFlag = False
            time.sleep(sleeptime)
            pageTemp = driver.page_source
            soup = BeautifulSoup(pageTemp,"html.parser")
            itemTemp = soup.find_all('div',class_="ContentItem-actions")
            for i in range(len(itemTemp)-4,len(itemTemp)):
                temp = str(itemTemp[i])
                like = convertNumStrToInt(re.findall(r'fill-rule="evenodd"></path></svg></span>赞同 (.*?)</button>',temp)[0])
                if like < minLikeTime:
                    flag = True
                    break
            if(flag):
                break
        #driver.execute_script('window.scrollBy(0,800000)')
        # Check out HTML entire page source codes for one page
        pageSourceThree = driver.page_source
    except Exception as e:
        flag = False
    return pageSourceThree,flag

# questionInfo : head + description + goodQuestionNum + questionCommentNum + questionAnswerNum
# questionAnswers[i] : content + likeNum + commentNum
def saveData(questionInfo,questionAnswers, savepath,url):
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.add_heading(questionInfo[0], level=1)
    doc.add_paragraph(questionInfo[1])
    doc.add_paragraph(url)

    doc_table = doc.add_table(rows=1, cols=3, style='Medium List 1 Accent 1')
    doc_table.rows[0].cells[0].text = '好问题数'
    doc_table.rows[0].cells[1].text = '评论数'
    doc_table.rows[0].cells[2].text = '答案数'
    new_row = doc_table.add_row()
    new_row.cells[0].text = str(questionInfo[2])
    new_row.cells[1].text = str(questionInfo[3])
    new_row.cells[2].text = str(questionInfo[4])

    doc.add_heading("热门回答", level=2)
    for i in range(math.ceil(len(questionAnswers) * percent)):
        doc.add_heading("回答"+str(i+1),level=3)
        doc.add_paragraph("点赞数" + str(questionAnswers[i][1]))
        doc.add_paragraph("评论数" + str(questionAnswers[i][2]))
        doc.add_paragraph("内容：\n"+questionAnswers[i][0])
        doc.add_paragraph("\n")

    doc.save(savepath)



if __name__ == "__main__":
    stage1 = ["https://www.zhihu.com/question/367716432","https://www.zhihu.com/question/367258958","https://www.zhihu.com/question/367530757","https://www.zhihu.com/question/367137465"]
    stage2 = ["https://www.zhihu.com/question/368177654","https://www.zhihu.com/question/368319975","https://www.zhihu.com/question/367426241","https://www.zhihu.com/question/367334538","https://www.zhihu.com/question/370089035","https://www.zhihu.com/question/367539714","https://www.zhihu.com/question/367539452","https://www.zhihu.com/question/367854449","https://www.zhihu.com/question/367279125","https://www.zhihu.com/question/367404487","https://www.zhihu.com/question/363965353","https://www.zhihu.com/question/367826502","https://www.zhihu.com/question/28634829","https://www.zhihu.com/question/370018105","https://www.zhihu.com/question/369927034","https://www.zhihu.com/question/367958339"]
    stage3 = ["https://www.zhihu.com/question/369333649","https://www.zhihu.com/question/368568986","https://www.zhihu.com/question/368477137","https://www.zhihu.com/question/370298868","https://www.zhihu.com/question/368248830","https://www.zhihu.com/question/369675155","https://www.zhihu.com/question/367967980","https://www.zhihu.com/question/370945152","https://www.zhihu.com/question/374246004","https://www.zhihu.com/question/370735342","https://www.zhihu.com/question/367319453","https://www.zhihu.com/question/368081151","https://www.zhihu.com/question/375257984","https://www.zhihu.com/question/375013479","https://www.zhihu.com/question/373548477","https://www.zhihu.com/question/371397753"]
    stage4 = ["https://www.zhihu.com/question/376738227","https://www.zhihu.com/question/380245498","https://www.zhihu.com/question/376144260","https://www.zhihu.com/question/380328444","https://www.zhihu.com/question/378560645","https://www.zhihu.com/question/378683936","https://www.zhihu.com/question/380113657","https://www.zhihu.com/question/397290358","https://www.zhihu.com/question/392963807","https://www.zhihu.com/question/393375587","https://www.zhihu.com/question/372862412","https://www.zhihu.com/question/379104604","https://www.zhihu.com/question/378134387","https://www.zhihu.com/question/370143181","https://www.zhihu.com/question/381454902","https://www.zhihu.com/question/378467357","https://www.zhihu.com/question/373467606","https://www.zhihu.com/question/385410772","https://www.zhihu.com/question/383478961","https://www.zhihu.com/question/389291230","https://www.zhihu.com/question/385360064","https://www.zhihu.com/question/375496794","https://www.zhihu.com/question/372824666","https://www.zhihu.com/question/379839194","https://www.zhihu.com/question/371062992","https://www.zhihu.com/question/379826353","https://www.zhihu.com/question/370156285","https://www.zhihu.com/question/374003316","https://www.zhihu.com/question/382938730","https://www.zhihu.com/question/371395246","https://www.zhihu.com/question/377416963/","https://www.zhihu.com/question/368791403","https://www.zhihu.com/question/381161861","https://www.zhihu.com/question/372364341","https://www.zhihu.com/question/395777438","https://www.zhihu.com/question/376512924"]
    main(stage1,stage2,stage3,stage4)
    # getData("https://www.zhihu.com/question/409954635")