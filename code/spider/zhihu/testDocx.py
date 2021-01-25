from docx import Document
from docx.oxml.ns import qn

Docx = Document()
Docx.styles['Normal'].font.name = u'宋体'
Docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
Docx.add_heading("这是一个一级标题",level=1)
Docx.add_paragraph("这是一个副级标题","Title")
A = Docx.add_paragraph("My name is aaa")
A.add_run("我学习的很快乐，啊哈哈哈哈哈，非常好 Good!!!")
Docx.add_heading("这是一个二级标题",level=2)
A = Docx.add_paragraph("这个是二级标题的内容呀")
B = A.add_run("二级标题里面的正文 继续添加！！！！！！！")
B.font.bold = True # 同时我要对这些正文进行加粗~~~~
B.font.size = (10)
Docx.add_heading("我爱学习Python以下就是python的logo呀",level=3)
doc_table = Docx.add_table(rows=1, cols=3, style='Medium List 1 Accent 1')

# 设置表头
doc_table.rows[0].cells[0].text = '姓名'
doc_table.rows[0].cells[1].text = '性别'
doc_table.rows[0].cells[2].text = '兴趣'

# 表体数据存储
date = (
    ("张三", "男","篮球"),
    ("李四", "男","乒乓球"),
    ("王舞", "女","羽毛球"),
)

# 添加数据，add_row()新增一行
for a,b,c in date:
    new_row = doc_table.add_row()
    new_row.cells[0].text = a
    new_row.cells[1].text = b
    new_row.cells[2].text = c

Docx.save("testDocx.docx")